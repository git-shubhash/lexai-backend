"""
Export route - generate PDF, DOCX, TXT exports
"""

import io
import logging
from flask import Blueprint, request, jsonify, send_file
from datetime import datetime
import traceback

logger = logging.getLogger(__name__)
export_bp = Blueprint('export', __name__)


def build_export_text(data: dict) -> str:
    """Build plain text export content"""
    lines = []
    lines.append("=" * 60)
    lines.append("LEGAL DOCUMENT ANALYSIS REPORT")
    lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("=" * 60)
    lines.append("")

    summary = data.get('summary', {})

    if summary.get('document_type'):
        lines.append(f"Document Type: {summary['document_type']}")
        lines.append("")

    if summary.get('short_summary'):
        lines.append("EXECUTIVE SUMMARY")
        lines.append("-" * 40)
        lines.append(summary['short_summary'])
        lines.append("")

    if summary.get('detailed_summary'):
        lines.append("DETAILED SUMMARY")
        lines.append("-" * 40)
        lines.append(summary['detailed_summary'])
        lines.append("")

    if summary.get('key_legal_points'):
        lines.append("KEY LEGAL POINTS")
        lines.append("-" * 40)
        for point in summary['key_legal_points']:
            lines.append(f"  • {point}")
        lines.append("")

    if summary.get('obligations'):
        lines.append("OBLIGATIONS")
        lines.append("-" * 40)
        for ob in summary['obligations']:
            lines.append(f"  • {ob}")
        lines.append("")

    if summary.get('important_deadlines'):
        lines.append("IMPORTANT DEADLINES")
        lines.append("-" * 40)
        for dl in summary['important_deadlines']:
            lines.append(f"  • {dl}")
        lines.append("")

    if summary.get('penalties'):
        lines.append("PENALTIES")
        lines.append("-" * 40)
        for p in summary['penalties']:
            lines.append(f"  • {p}")
        lines.append("")

    if summary.get('legal_risks'):
        lines.append("RISK ANALYSIS")
        lines.append("-" * 40)
        for risk in summary['legal_risks']:
            level = risk.get('level', 'unknown').upper()
            lines.append(f"  [{level}] {risk.get('risk', '')}")
            if risk.get('explanation'):
                lines.append(f"         {risk['explanation']}")
        lines.append("")

    clauses = summary.get('clauses', {})
    if any(clauses.values()):
        lines.append("EXTRACTED CLAUSES")
        lines.append("-" * 40)
        for clause_name, clause_text in clauses.items():
            if clause_text:
                lines.append(f"\n{clause_name.upper().replace('_', ' ')}:")
                lines.append(clause_text)
        lines.append("")

    lines.append("=" * 60)
    lines.append("END OF REPORT")
    return "\n".join(lines)


@export_bp.route('/export/txt', methods=['POST'])
def export_txt():
    """Export analysis as TXT"""
    data = request.get_json(silent=True)
    if not data:
        return jsonify({'error': 'JSON body required'}), 400

    try:
        content = build_export_text(data)
        buf = io.BytesIO(content.encode('utf-8'))
        buf.seek(0)
        filename = data.get('filename', 'legal_analysis').replace(' ', '_')

        return send_file(
            buf,
            mimetype='text/plain',
            as_attachment=True,
            download_name=f"{filename}_analysis.txt"
        )
    except Exception as e:
        logger.error(f"TXT export error: {e}")
        return jsonify({'error': str(e)}), 500


@export_bp.route('/export/docx', methods=['POST'])
def export_docx():
    """Export analysis as DOCX"""
    data = request.get_json(silent=True)
    if not data:
        return jsonify({'error': 'JSON body required'}), 400

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        summary = data.get('summary', {})

        # Title
        title = doc.add_heading('Legal Document Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        if summary.get('document_type'):
            doc.add_paragraph(f"Document Type: {summary['document_type']}")

        if summary.get('short_summary'):
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(summary['short_summary'])

        if summary.get('detailed_summary'):
            doc.add_heading('Detailed Summary', level=1)
            doc.add_paragraph(summary['detailed_summary'])

        if summary.get('key_legal_points'):
            doc.add_heading('Key Legal Points', level=1)
            for point in summary['key_legal_points']:
                doc.add_paragraph(point, style='List Bullet')

        if summary.get('obligations'):
            doc.add_heading('Obligations', level=1)
            for ob in summary['obligations']:
                doc.add_paragraph(ob, style='List Bullet')

        if summary.get('legal_risks'):
            doc.add_heading('Risk Analysis', level=1)
            for risk in summary['legal_risks']:
                p = doc.add_paragraph()
                p.add_run(f"[{risk.get('level', '').upper()}] ").bold = True
                p.add_run(risk.get('risk', ''))

        clauses = summary.get('clauses', {})
        if any(clauses.values()):
            doc.add_heading('Extracted Clauses', level=1)
            for clause_name, clause_text in clauses.items():
                if clause_text:
                    doc.add_heading(clause_name.replace('_', ' ').title(), level=2)
                    doc.add_paragraph(clause_text)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        filename = data.get('filename', 'legal_analysis').replace(' ', '_')

        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"{filename}_analysis.docx"
        )
    except Exception as e:
        logger.error(f"DOCX export error: {e}")
        return jsonify({'error': str(e)}), 500


@export_bp.route('/export/pdf', methods=['POST'])
def export_pdf():
    """Export analysis as PDF"""
    data = request.get_json(silent=True)
    if not data:
        return jsonify({'error': 'JSON body required'}), 400

    try:
        from fpdf import FPDF
        
        class PDF(FPDF):
            def header(self):
                self.set_font('helvetica', 'B', 15)
                self.cell(0, 10, 'Legal Document Analysis Report', 0, 1, 'C')
                self.ln(5)

            def footer(self):
                self.set_y(-15)
                self.set_font('helvetica', 'I', 8)
                self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

        def clean_text(text):
            if not text: return ""
            # Replace common special characters that crash fpdf2 with standard ones
            replacements = {
                '\u2013': '-', '\u2014': '-', '\u2018': "'", '\u2019': "'",
                '\u201c': '"', '\u201d': '"', '\u2022': '*', '\u00a7': 'Sect.',
                '\u2122': '(TM)', '\u00ae': '(R)', '\u00a9': '(C)'
            }
            for char, rep in replacements.items():
                text = text.replace(char, rep)
            # Encode and decode as latin-1 to strip any other unknown characters
            return text.encode('latin-1', 'replace').decode('latin-1')

        pdf = PDF()
        pdf.add_page()
        pdf.set_font("helvetica", size=10)
        
        summary = data.get('summary', {})
        
        pdf.set_font("helvetica", 'B', 10)
        pdf.cell(0, 10, clean_text(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"), 0, 1)
        
        if summary.get('document_type'):
            pdf.cell(0, 10, clean_text(f"Document Type: {summary['document_type']}"), 0, 1)
        
        pdf.ln(5)

        # Executive Summary
        if summary.get('short_summary'):
            pdf.set_font("helvetica", 'B', 12)
            pdf.cell(0, 10, "Executive Summary", 0, 1)
            pdf.set_font("helvetica", size=10)
            pdf.multi_cell(0, 5, clean_text(summary['short_summary']))
            pdf.ln(5)

        # Detailed Summary
        if summary.get('detailed_summary'):
            pdf.set_font("helvetica", 'B', 12)
            pdf.cell(0, 10, "Detailed Summary", 0, 1)
            pdf.set_font("helvetica", size=10)
            pdf.multi_cell(0, 5, clean_text(summary['detailed_summary']))
            pdf.ln(5)

        # Key Legal Points
        if summary.get('key_legal_points'):
            pdf.set_font("helvetica", 'B', 12)
            pdf.cell(0, 10, "Key Legal Points", 0, 1)
            pdf.set_font("helvetica", size=10)
            for point in summary['key_legal_points']:
                pdf.multi_cell(0, 5, clean_text(f"• {point}"))
            pdf.ln(5)

        # Obligations
        if summary.get('obligations'):
            pdf.set_font("helvetica", 'B', 12)
            pdf.cell(0, 10, "Obligations", 0, 1)
            pdf.set_font("helvetica", size=10)
            for ob in summary['obligations']:
                pdf.multi_cell(0, 5, clean_text(f"• {ob}"))
            pdf.ln(5)

        # Risk Analysis
        if summary.get('legal_risks'):
            pdf.set_font("helvetica", 'B', 12)
            pdf.cell(0, 10, "Risk Analysis", 0, 1)
            pdf.set_font("helvetica", size=10)
            for risk in summary['legal_risks']:
                level = risk.get('level', '').upper()
                pdf.set_font("helvetica", 'B', 10)
                pdf.cell(0, 5, clean_text(f"[{level}] {risk.get('risk', '')}"), 0, 1)
                pdf.set_font("helvetica", size=10)
                pdf.multi_cell(0, 5, clean_text(risk.get('explanation', '')))
                pdf.ln(2)
            pdf.ln(5)

        # Clauses
        clauses = summary.get('clauses', {})
        if any(clauses.values()):
            pdf.set_font("helvetica", 'B', 12)
            pdf.cell(0, 10, "Extracted Clauses", 0, 1)
            for name, text in clauses.items():
                if text:
                    pdf.set_font("helvetica", 'B', 11)
                    pdf.cell(0, 8, clean_text(name.replace('_', ' ').title()), 0, 1)
                    pdf.set_font("helvetica", size=10)
                    pdf.multi_cell(0, 5, clean_text(text))
                    pdf.ln(3)

        buf = io.BytesIO(pdf.output())
        buf.seek(0)
        filename = data.get('filename', 'legal_analysis').replace(' ', '_')

        return send_file(
            buf,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f"{filename}_analysis.pdf"
        )
    except Exception as e:
        logger.error(f"PDF export error: {traceback.format_exc() if 'traceback' in globals() else e}")
        return jsonify({'error': str(e), 'details': 'Ensure all special characters are Latin-1 compatible'}), 500
