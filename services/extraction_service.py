"""
Document extraction service
Handles PDF, DOCX, TXT, and image OCR extraction
"""

import io
import os
import logging
import tempfile
from typing import Tuple

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> Tuple[str, str]:
    """Extract text from PDF using pdfplumber with PyPDF2 fallback"""
    text = ""
    method = "pdfplumber"

    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages_text = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    pages_text.append(page_text)
            text = "\n\n".join(pages_text)

        if text.strip():
            logger.info(f"PDF extracted with pdfplumber: {len(text)} chars")
            return text, method
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}")

    # Fallback to PyPDF2
    try:
        import PyPDF2
        method = "PyPDF2"
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages_text = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                pages_text.append(page_text)
        text = "\n\n".join(pages_text)

        if text.strip():
            logger.info(f"PDF extracted with PyPDF2: {len(text)} chars")
            return text, method
    except Exception as e:
        logger.warning(f"PyPDF2 failed: {e}")

    # Try OCR if text extraction failed (scanned PDF)
    if not text.strip():
        logger.info("Attempting OCR for scanned PDF")
        text, _ = extract_text_via_ocr(file_bytes, is_pdf=True)
        method = "ocr"

    return text, method


def extract_text_from_docx(file_bytes: bytes) -> Tuple[str, str]:
    """Extract text from DOCX file"""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        text = "\n\n".join(paragraphs)
        logger.info(f"DOCX extracted: {len(text)} chars")
        return text, "python-docx"
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise e


def extract_text_from_txt(file_bytes: bytes) -> Tuple[str, str]:
    """Extract text from TXT file"""
    try:
        # Try UTF-8 first, then latin-1
        try:
            text = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            text = file_bytes.decode('latin-1')
        logger.info(f"TXT extracted: {len(text)} chars")
        return text, "plain-text"
    except Exception as e:
        logger.error(f"TXT extraction failed: {e}")
        raise e


def extract_text_via_ocr(file_bytes: bytes, is_pdf: bool = False) -> Tuple[str, str]:
    """Extract text from images or scanned PDFs using Tesseract OCR"""
    try:
        import pytesseract
        from PIL import Image

        # Configure tesseract path for Windows
        if os.name == 'nt':
            tesseract_paths = [
                r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            ]
            for path in tesseract_paths:
                if os.path.exists(path):
                    pytesseract.pytesseract.tesseract_cmd = path
                    break

        texts = []

        if is_pdf:
            try:
                from pdf2image import convert_from_bytes
                images = convert_from_bytes(file_bytes, dpi=200)
                for img in images:
                    text = pytesseract.image_to_string(img, lang='eng')
                    if text.strip():
                        texts.append(text)
            except Exception as e:
                logger.warning(f"pdf2image failed: {e}")
        else:
            image = Image.open(io.BytesIO(file_bytes))
            text = pytesseract.image_to_string(image, lang='eng')
            if text.strip():
                texts.append(text)

        result = "\n\n".join(texts)
        logger.info(f"OCR extracted: {len(result)} chars")
        return result, "tesseract-ocr"

    except Exception as e:
        logger.error(f"OCR extraction failed: {e}")
        return "", "ocr-failed"


def extract_text(file_bytes: bytes, filename: str) -> dict:
    """Main extraction dispatcher based on file type"""
    filename_lower = filename.lower()
    text = ""
    method = "unknown"

    try:
        if filename_lower.endswith('.pdf'):
            text, method = extract_text_from_pdf(file_bytes)
        elif filename_lower.endswith('.docx'):
            text, method = extract_text_from_docx(file_bytes)
        elif filename_lower.endswith('.txt'):
            text, method = extract_text_from_txt(file_bytes)
        elif filename_lower.endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif')):
            text, method = extract_text_via_ocr(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {filename}")

        word_count = len(text.split()) if text else 0
        char_count = len(text) if text else 0

        return {
            'success': True,
            'text': text,
            'method': method,
            'word_count': word_count,
            'char_count': char_count,
            'is_empty': not bool(text.strip())
        }

    except Exception as e:
        logger.error(f"Extraction failed for {filename}: {e}")
        return {
            'success': False,
            'text': '',
            'method': method,
            'error': str(e),
            'word_count': 0,
            'char_count': 0,
            'is_empty': True
        }
